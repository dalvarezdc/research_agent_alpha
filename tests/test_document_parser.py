#!/usr/bin/env python3
"""Tests for the document_parser service.

Fixtures are generated in-test: docx via python-docx, PDFs via a tiny
hand-built PDF writer (no system libraries required), text/rtf inline.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from document_parser import (
    ParseStatus,
    parse_document,
    supported_extensions,
)
from document_parser.backends.base import (
    looks_like_heading,
    render_list,
    render_table,
    split_paragraphs,
)


# ── Minimal PDF fixture helper (no external deps / system libs) ───────────────


def _make_pdf(pages: list[list[str]]) -> bytes:
    """Build a minimal multi-page PDF whose pages contain the given text lines.

    An empty line list for a page produces a page with no text (simulating a
    scanned/image-only page).
    """
    objs: list[bytes] = []
    catalog_idx = 1
    pages_idx = 2

    page_obj_indices: list[int] = []
    content_objs: list[bytes] = []

    # Reserve object numbering: catalog(1), pages(2), then per page: page + content.
    next_idx = 3
    for lines in pages:
        page_idx = next_idx
        content_idx = next_idx + 1
        next_idx += 2
        page_obj_indices.append(page_idx)

        content = "BT /F1 12 Tf 72 720 Td 14 TL\n"
        for i, line in enumerate(lines):
            esc = line.replace("(", r"\(").replace(")", r"\)")
            content += (f"({esc}) Tj\n" if i == 0 else f"T* ({esc}) Tj\n")
        content += "ET"
        cb = content.encode("latin-1")
        content_objs.append(
            b"<< /Length " + str(len(cb)).encode() + b" >>\nstream\n" + cb + b"\nendstream"
        )

    font_idx = next_idx

    # Assemble objects in index order.
    numbered: dict[int, bytes] = {}
    numbered[catalog_idx] = b"<< /Type /Catalog /Pages 2 0 R >>"
    kids = b" ".join(f"{i} 0 R".encode() for i in page_obj_indices)
    numbered[pages_idx] = (
        b"<< /Type /Pages /Kids [" + kids + b"] /Count "
        + str(len(page_obj_indices)).encode() + b" >>"
    )
    for offset, page_idx in enumerate(page_obj_indices):
        content_idx = page_idx + 1
        numbered[page_idx] = (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents "
            + f"{content_idx} 0 R".encode()
            + b" /Resources << /Font << /F1 "
            + f"{font_idx} 0 R".encode()
            + b" >> >> >>"
        )
        numbered[content_idx] = content_objs[offset]
    numbered[font_idx] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    total = font_idx
    pdf = b"%PDF-1.4\n"
    offsets: list[int] = []
    for i in range(1, total + 1):
        offsets.append(len(pdf))
        pdf += f"{i} 0 obj\n".encode() + numbered[i] + b"\nendobj\n"
    xref_pos = len(pdf)
    pdf += b"xref\n0 " + str(total + 1).encode() + b"\n"
    pdf += b"0000000000 65535 f \n"
    for off in offsets:
        pdf += f"{off:010d} 00000 n \n".encode()
    pdf += (
        b"trailer\n<< /Size " + str(total + 1).encode()
        + b" /Root 1 0 R >>\nstartxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    )
    return pdf


@pytest.fixture()
def text_pdf(tmp_path) -> Path:
    p = tmp_path / "report.pdf"
    p.write_bytes(
        _make_pdf([["MEDICAL REPORT", "Patient shows normal levels.", "Follow up soon."]])
    )
    return p


@pytest.fixture()
def multipage_pdf(tmp_path) -> Path:
    p = tmp_path / "multi.pdf"
    p.write_bytes(_make_pdf([["Page one text."], ["Page two text."]]))
    return p


@pytest.fixture()
def empty_pdf(tmp_path) -> Path:
    """A PDF whose only page has no text (image-only simulation)."""
    p = tmp_path / "empty.pdf"
    p.write_bytes(_make_pdf([[]]))
    return p


@pytest.fixture()
def partial_pdf(tmp_path) -> Path:
    """One page with text, one page without."""
    p = tmp_path / "partial.pdf"
    p.write_bytes(_make_pdf([["Has text here."], []]))
    return p


# ── base.py helper tests ─────────────────────────────────────────────────────


def test_render_table_basic():
    md = render_table([["A", "B"], ["1", "2"]])
    assert md == "| A | B |\n| --- | --- |\n| 1 | 2 |"


def test_render_table_pads_short_rows_and_escapes_pipes():
    md = render_table([["A", "B"], ["x|y"]])
    lines = md.splitlines()
    assert lines[0] == "| A | B |"
    assert lines[2] == "| x\\|y |  |"  # padded + escaped


def test_render_table_empty_returns_blank():
    assert render_table([]) == ""
    assert render_table([[None]]) == "|  |\n| --- |"


def test_render_list_ordered_and_unordered():
    assert render_list(["a", "b"]) == "- a\n- b"
    assert render_list(["a", "b"], ordered=True) == "1. a\n2. b"
    assert render_list(["", "  "]) == ""  # blank items skipped


def test_split_paragraphs():
    text = "Line one\nline two\n\nSecond para\n\n\nThird"
    assert split_paragraphs(text) == ["Line one line two", "Second para", "Third"]


@pytest.mark.parametrize(
    "line,expected",
    [
        ("MEDICAL REPORT", True),
        ("SUMMARY", True),
        ("This is a normal sentence.", False),
        ("lowercase heading", False),
        ("A", True),
        ("This heading is way too long to be considered a heading at all here", False),
        ("", False),
        ("Ends with colon:", False),
    ],
)
def test_looks_like_heading(line, expected):
    assert looks_like_heading(line) is expected


# ── PDF backend ──────────────────────────────────────────────────────────────


def test_parse_text_pdf_success(text_pdf):
    result = parse_document(text_pdf)
    assert result.status is ParseStatus.SUCCESS
    assert "Patient shows normal levels." in result.markdown
    assert "## MEDICAL REPORT" in result.markdown  # heading heuristic
    assert result.metadata.page_count == 1
    assert result.metadata.backend == "pdf"
    assert result.metadata.char_count > 0


def test_parse_multipage_pdf_has_page_markers(multipage_pdf):
    result = parse_document(multipage_pdf)
    assert result.status is ParseStatus.SUCCESS
    assert "<!-- page 1 -->" in result.markdown
    assert "<!-- page 2 -->" in result.markdown
    assert result.metadata.page_count == 2


def test_parse_empty_pdf_is_failed_with_warning(empty_pdf):
    result = parse_document(empty_pdf)
    assert result.status is ParseStatus.FAILED
    assert result.warnings
    assert any("page 1" in str(w) for w in result.warnings)


def test_parse_partial_pdf(partial_pdf):
    result = parse_document(partial_pdf)
    assert result.status is ParseStatus.PARTIAL
    assert "Has text here." in result.markdown
    assert any("page 2" in str(w) for w in result.warnings)


# ── docx backend ─────────────────────────────────────────────────────────────


@pytest.fixture()
def sample_docx(tmp_path) -> Path:
    import docx

    document = docx.Document()
    document.add_heading("Lab Results", level=1)
    document.add_paragraph("The patient is stable.")
    document.add_paragraph("First item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Test"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Vitamin D"
    table.rows[1].cells[1].text = "18 ng/mL"
    p = tmp_path / "labs.docx"
    document.save(str(p))
    return p


def test_parse_docx_headings_paragraphs_tables(sample_docx):
    result = parse_document(sample_docx)
    assert result.status is ParseStatus.SUCCESS
    assert "# Lab Results" in result.markdown
    assert "The patient is stable." in result.markdown
    assert "- First item" in result.markdown
    assert "| Test | Value |" in result.markdown
    assert "| Vitamin D | 18 ng/mL |" in result.markdown
    assert result.metadata.backend == "docx"


# ── text backend ─────────────────────────────────────────────────────────────


def test_parse_txt(tmp_path):
    p = tmp_path / "note.txt"
    p.write_text("Para one line1\nline2\n\nPara two", encoding="utf-8")
    result = parse_document(p)
    assert result.status is ParseStatus.SUCCESS
    assert result.markdown == "Para one line1 line2\n\nPara two"
    assert result.metadata.file_format == "txt"


def test_parse_md_passthrough(tmp_path):
    p = tmp_path / "doc.md"
    p.write_text("# Title\n\nSome **bold** text.", encoding="utf-8")
    result = parse_document(p)
    assert result.status is ParseStatus.SUCCESS
    assert result.markdown == "# Title\n\nSome **bold** text."


def test_parse_empty_txt_is_failed(tmp_path):
    p = tmp_path / "blank.txt"
    p.write_text("   \n\n  ", encoding="utf-8")
    result = parse_document(p)
    assert result.status is ParseStatus.FAILED
    assert result.warnings


# ── rtf backend (optional dep, installed via parsing-extras) ──────────────────


def test_parse_rtf(tmp_path):
    pytest.importorskip("striprtf")
    p = tmp_path / "note.rtf"
    p.write_text(
        r"{\rtf1\ansi\deff0 {\fonttbl{\f0 Times;}}\f0\fs24 Hello RTF world.\par}",
        encoding="utf-8",
    )
    result = parse_document(p)
    assert result.status is ParseStatus.SUCCESS
    assert "Hello RTF world." in result.markdown
    assert result.metadata.backend == "rtf"


# ── dispatcher behavior ──────────────────────────────────────────────────────


def test_unsupported_extension_is_failed(tmp_path):
    p = tmp_path / "data.xyz"
    p.write_text("whatever", encoding="utf-8")
    result = parse_document(p)
    assert result.status is ParseStatus.FAILED
    assert any("Unsupported format" in str(w) for w in result.warnings)


def test_missing_path_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        parse_document(tmp_path / "does_not_exist.pdf")


def test_content_type_override(tmp_path):
    """A misnamed file can be parsed via an explicit format override."""
    p = tmp_path / "note.data"
    p.write_text("hello override", encoding="utf-8")
    result = parse_document(p, content_type="txt")
    assert result.status is ParseStatus.SUCCESS
    assert "hello override" in result.markdown


def test_missing_optional_dependency_reports_install_hint(tmp_path, monkeypatch):
    """When a backend's optional dep is unavailable, dispatch returns a helpful
    FAILED result instead of crashing."""
    from document_parser import service

    backend = service._REGISTRY[".rtf"]
    monkeypatch.setattr(backend, "is_available", lambda: False)

    p = tmp_path / "note.rtf"
    p.write_text("irrelevant", encoding="utf-8")
    result = parse_document(p)
    assert result.status is ParseStatus.FAILED
    assert any("requires" in str(w) for w in result.warnings)


def test_backend_exception_is_caught(tmp_path, monkeypatch):
    """An unexpected error inside a backend becomes a FAILED result, not a crash."""
    from document_parser import service

    backend = service._REGISTRY[".txt"]

    def _boom(_path):
        raise RuntimeError("simulated backend explosion")

    monkeypatch.setattr(backend, "parse", _boom)

    p = tmp_path / "note.txt"
    p.write_text("hello", encoding="utf-8")
    result = parse_document(p)
    assert result.status is ParseStatus.FAILED
    assert any("Failed to parse" in str(w) for w in result.warnings)


def test_supported_extensions_includes_priority_formats():
    exts = supported_extensions()
    for ext in (".pdf", ".docx", ".txt", ".md"):
        assert ext in exts


def test_parse_result_ok_property(text_pdf):
    assert parse_document(text_pdf).ok is True


# ── CLI ──────────────────────────────────────────────────────────────────────


def test_cli_to_stdout(text_pdf, capsys):
    from document_parser.__main__ import main

    exit_code = main([str(text_pdf)])
    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Patient shows normal levels." in captured.out


def test_cli_to_output_file(text_pdf, tmp_path):
    from document_parser.__main__ import main

    out = tmp_path / "out.md"
    exit_code = main([str(text_pdf), "-o", str(out)])
    assert exit_code == 0
    assert "Patient shows normal levels." in out.read_text(encoding="utf-8")


def test_cli_missing_file_returns_error(tmp_path, capsys):
    from document_parser.__main__ import main

    exit_code = main([str(tmp_path / "nope.pdf")])
    captured = capsys.readouterr()
    assert exit_code == 1
    assert "error:" in captured.err


def test_cli_failed_status_exit_code(tmp_path, capsys):
    from document_parser.__main__ import main

    p = tmp_path / "data.xyz"
    p.write_text("x", encoding="utf-8")
    exit_code = main([str(p)])
    assert exit_code == 1


# ── Additional backend unit tests (targeted coverage) ────────────────────────


def test_docx_render_paragraph_variants():
    """Exercise title, heading-level, numbered/bulleted list, and empty paths."""
    from document_parser.backends.docx_backend import DocxBackend, _heading_level

    class _Style:
        def __init__(self, name):
            self.name = name

    class _Para:
        def __init__(self, text, style_name):
            self.text = text
            self.style = _Style(style_name)

    render = DocxBackend._render_paragraph
    assert render(_Para("", "Normal")) == ""
    assert render(_Para("Big Title", "Title")) == "# Big Title"
    assert render(_Para("Sub", "Heading 2")) == "## Sub"
    assert render(_Para("Step", "List Number")) == "1. Step"
    assert render(_Para("Bullet", "List Bullet")) == "- Bullet"
    assert render(_Para("Plain", "Normal")) == "Plain"

    assert _heading_level("heading 3") == 3
    assert _heading_level("heading") == 2  # no digit → default
    assert _heading_level("heading 99") == 6  # clamped


def test_pdf_render_text_blocks_groups_and_headings():
    from document_parser.backends.pdf_backend import PdfBackend

    text = "MEDICAL REPORT\nLine a\nline b\n\nSUMMARY\nStable patient."
    blocks = PdfBackend._render_text_blocks(text)
    assert blocks == [
        "## MEDICAL REPORT",
        "Line a line b",
        "## SUMMARY",
        "Stable patient.",
    ]


def test_pdf_pypdf_fallback_used_when_pdfplumber_empty(tmp_path, monkeypatch):
    """If pdfplumber yields no text, pypdf fallback text is used."""
    from document_parser.backends import pdf_backend

    p = tmp_path / "fb.pdf"
    p.write_bytes(_make_pdf([["Real content via pypdf"]]))

    # Force pdfplumber page text to empty so the fallback path runs.
    import pdfplumber

    monkeypatch.setattr(
        pdfplumber.page.Page, "extract_text", lambda self, *a, **k: "", raising=True
    )
    monkeypatch.setattr(
        pdfplumber.page.Page, "extract_tables", lambda self, *a, **k: [], raising=True
    )

    result = parse_document(p)
    # pypdf should recover the text → SUCCESS with the content present.
    assert result.status is ParseStatus.SUCCESS
    assert "Real content via pypdf" in result.markdown


def test_doc_backend_unavailable_reports_hint(tmp_path, monkeypatch):
    """Legacy .doc without textract yields a FAILED result with install hint."""
    from document_parser import service

    backend = service._REGISTRY[".doc"]
    monkeypatch.setattr(backend, "is_available", lambda: False)

    p = tmp_path / "legacy.doc"
    p.write_bytes(b"\xd0\xcf\x11\xe0stub")  # OLE-ish header bytes
    result = parse_document(p)
    assert result.status is ParseStatus.FAILED
    assert any("textract" in str(w) for w in result.warnings)


def test_cli_output_file_reports_status_on_stderr(text_pdf, tmp_path, capsys):
    from document_parser.__main__ import main

    out = tmp_path / "o.md"
    main([str(text_pdf), "-o", str(out)])
    captured = capsys.readouterr()
    assert "success" in captured.err


def test_pypdf_fallback_edge_cases(tmp_path):
    """Out-of-range page index and unreadable file both return empty string."""
    from document_parser.backends.pdf_backend import PdfBackend

    p = tmp_path / "one.pdf"
    p.write_bytes(_make_pdf([["only page"]]))

    # Page index beyond the document → "" (final return).
    assert PdfBackend._pypdf_fallback(p, 99) == ""

    # Corrupt/unreadable file → exception path → "".
    bad = tmp_path / "bad.pdf"
    bad.write_bytes(b"not a real pdf")
    assert PdfBackend._pypdf_fallback(bad, 1) == ""


def test_pdf_table_rendering(tmp_path, monkeypatch):
    """A page exposing a table grid renders a GFM table in the markdown."""
    from document_parser.backends import pdf_backend

    p = tmp_path / "tbl.pdf"
    p.write_bytes(_make_pdf([["Body text."]]))

    import pdfplumber

    monkeypatch.setattr(
        pdfplumber.page.Page,
        "extract_tables",
        lambda self, *a, **k: [[["Test", "Value"], ["Vitamin D", "18"]]],
        raising=True,
    )

    result = parse_document(p)
    assert "| Test | Value |" in result.markdown
    assert "| Vitamin D | 18 |" in result.markdown


def test_pdf_table_extraction_error_is_swallowed(tmp_path, monkeypatch):
    """If table extraction raises, parsing still succeeds on the text."""
    from document_parser.backends import pdf_backend

    p = tmp_path / "tblerr.pdf"
    p.write_bytes(_make_pdf([["Body text here."]]))

    import pdfplumber

    def _raise(self, *a, **k):
        raise RuntimeError("table boom")

    monkeypatch.setattr(pdfplumber.page.Page, "extract_tables", _raise, raising=True)

    result = parse_document(p)
    assert result.status is ParseStatus.SUCCESS
    assert "Body text here." in result.markdown
