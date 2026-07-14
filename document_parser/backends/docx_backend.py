"""Microsoft Word (.docx) backend using python-docx.

Maps Word heading styles to markdown headings, paragraphs to text, list
paragraphs to markdown lists, and tables to GFM tables. Core backend (python-docx
is a base dependency) but the import is still guarded so a broken install
degrades gracefully rather than crashing at import time.
"""

from __future__ import annotations

from pathlib import Path

from ..models import DocumentMetadata, ParseResult, ParseStatus
from .base import ParserBackend, render_table

try:
    import docx  # type: ignore

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - python-docx is a base dependency
    docx = None  # type: ignore
    _DOCX_AVAILABLE = False


class DocxBackend(ParserBackend):
    name = "docx"
    extensions = (".docx",)
    required_import = "docx"

    def is_available(self) -> bool:
        return _DOCX_AVAILABLE

    def parse(self, path: Path) -> ParseResult:
        document = docx.Document(str(path))
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            block = self._render_paragraph(paragraph)
            if block:
                blocks.append(block)

        for table in document.tables:
            rendered = self._render_table(table)
            if rendered:
                blocks.append(rendered)

        markdown = "\n\n".join(blocks).strip()
        status = ParseStatus.SUCCESS if markdown else ParseStatus.FAILED
        result = ParseResult(
            markdown=markdown,
            status=status,
            metadata=DocumentMetadata(
                source_path=str(path),
                file_format="docx",
                backend=self.name,
                char_count=len(markdown),
            ),
        )
        if status is ParseStatus.FAILED:
            result.add_warning("Document contained no extractable text or tables.")
        return result

    @staticmethod
    def _render_paragraph(paragraph: docx.text.paragraph.Paragraph) -> str:
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = (getattr(paragraph.style, "name", "") or "").lower()

        if style_name.startswith("heading"):
            level = _heading_level(style_name)
            return f"{'#' * level} {text}"
        if style_name.startswith("title"):
            return f"# {text}"
        if "list" in style_name:
            # Numbered vs. bulleted: python-docx style names contain "number".
            prefix = "1." if "number" in style_name else "-"
            return f"{prefix} {text}"
        return text

    @staticmethod
    def _render_table(table: docx.table.Table) -> str:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        return render_table(rows)


def _heading_level(style_name: str) -> int:
    """Extract a 1-6 heading level from a style name like 'heading 2'."""
    for token in style_name.split():
        if token.isdigit():
            return min(max(int(token), 1), 6)
    return 2
